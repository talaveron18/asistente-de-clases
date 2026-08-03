from __future__ import annotations

import json
from pathlib import Path
from typing import Callable


class ErrorExtraccionDocumento(RuntimeError):
    pass


def extraer_documento(ruta: str, callback: Callable[[str, float], None] | None = None) -> dict:
    path = Path(ruta)
    if not path.is_file():
        raise FileNotFoundError(ruta)

    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extraer_pdf(path, callback)
    if ext == ".docx":
        return _extraer_docx(path, callback)
    if ext in {".txt", ".md"}:
        texto = path.read_text(encoding="utf-8", errors="replace")
        return {
            "tipo": ext[1:],
            "paginas": [{"pagina": 1, "texto": texto}],
            "caracteres": len(texto),
            "requiere_ocr": False,
        }
    raise ErrorExtraccionDocumento(f"Formato no compatible para extracción: {ext}")


def _extraer_pdf(path: Path, callback=None) -> dict:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ErrorExtraccionDocumento("Falta la dependencia pypdf.") from exc

    reader = PdfReader(str(path))
    paginas = []
    total = max(1, len(reader.pages))
    caracteres = 0
    paginas_vacias = 0

    for i, page in enumerate(reader.pages, 1):
        if callback:
            callback(f"Leyendo página {i} de {total}...", i / total)
        try:
            texto = page.extract_text() or ""
        except Exception:
            texto = ""
        texto = texto.strip()
        if len(texto) < 20:
            paginas_vacias += 1
        caracteres += len(texto)
        paginas.append({"pagina": i, "texto": texto})

    ratio_vacias = paginas_vacias / total
    requiere_ocr = caracteres < 200 or ratio_vacias > 0.8
    return {
        "tipo": "pdf",
        "paginas": paginas,
        "caracteres": caracteres,
        "requiere_ocr": requiere_ocr,
        "total_paginas": total,
    }


def _extraer_docx(path: Path, callback=None) -> dict:
    try:
        from docx import Document
    except ImportError as exc:
        raise ErrorExtraccionDocumento("Falta la dependencia python-docx.") from exc

    if callback:
        callback("Leyendo documento Word...", 0.2)
    doc = Document(str(path))
    bloques = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto:
            bloques.append(texto)
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                bloques.append(" | ".join(celdas))
    texto = "\n".join(bloques)
    if callback:
        callback("Documento Word procesado.", 1.0)
    return {
        "tipo": "docx",
        "paginas": [{"pagina": 1, "texto": texto}],
        "caracteres": len(texto),
        "requiere_ocr": False,
    }


def guardar_extraccion(resultado: dict, destino: str) -> None:
    Path(destino).write_text(
        json.dumps(resultado, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
