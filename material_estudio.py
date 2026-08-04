from __future__ import annotations

import csv
import json
import re
from pathlib import Path


def generar_material_estudio(carpeta: str | Path) -> dict:
    carpeta = Path(carpeta)
    datos = json.loads((carpeta / "pipeline_clase.json").read_text(encoding="utf-8"))
    flashcards = []
    preguntas = []

    for bloque in datos.get("bloques", []):
        titulo = bloque.get("titulo", "Bloque")
        resumen = bloque.get("resumen", "").strip()
        claves = bloque.get("palabras_clave", [])
        if resumen:
            flashcards.append((f"Resume {titulo}", resumen, bloque.get("inicio", "")))
        for clave in claves[:5]:
            respuesta = _frase_relevante(bloque.get("texto", ""), clave)
            if respuesta:
                flashcards.append((f"¿Qué se explicó sobre {clave}?", respuesta, bloque.get("inicio", "")))
        for pregunta in bloque.get("preguntas", []):
            limpia = re.sub(r"^\[[^\]]+\]\s*", "", pregunta).strip()
            if limpia:
                preguntas.append((limpia, bloque.get("inicio", ""), titulo))
        if not bloque.get("preguntas"):
            preguntas.append((f"Explica los puntos esenciales de {titulo}.", bloque.get("inicio", ""), titulo))

    vistos = set()
    flashcards_unicas = []
    for frente, dorso, minuto in flashcards:
        clave = (frente.casefold(), dorso.casefold())
        if clave not in vistos:
            vistos.add(clave)
            flashcards_unicas.append((frente, dorso, minuto))

    with (carpeta / "flashcards_argos.tsv").open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f, delimiter="\t")
        writer.writerow(["Frente", "Dorso", "Minuto"])
        writer.writerows(flashcards_unicas)

    lineas = [f"# Preguntas de repaso · {datos.get('titulo', '')}", ""]
    for i, (pregunta, minuto, bloque) in enumerate(preguntas, 1):
        lineas += [f"{i}. **{pregunta}**", f"   - Bloque: {bloque}", f"   - Referencia: {minuto}", ""]
    (carpeta / "preguntas_repaso.md").write_text("\n".join(lineas), encoding="utf-8")

    hoja = [
        f"# Hoja de repaso rápido · {datos.get('titulo', '')}", "",
        f"**Materia:** {datos.get('materia', '')}", "",
        "## Conceptos dominantes", "",
        ", ".join(datos.get("palabras_clave_globales", [])) or "—", "",
        "## Avisos de examen", "",
    ]
    hoja += [f"- {x}" for x in datos.get("avisos_examen", [])] or ["- No se detectaron avisos explícitos."]
    hoja += ["", "## Bloques esenciales", ""]
    for bloque in datos.get("bloques", []):
        hoja += [
            f"### {bloque.get('numero')}. {bloque.get('titulo')}",
            f"**Minuto:** {bloque.get('inicio')}–{bloque.get('fin')}", "",
            bloque.get("resumen", ""), "",
        ]
    (carpeta / "repaso_rapido.md").write_text("\n".join(hoja), encoding="utf-8")

    docx_generado = _generar_docx(carpeta, datos, preguntas, flashcards_unicas)
    archivos = ["flashcards_argos.tsv", "preguntas_repaso.md", "repaso_rapido.md"]
    if docx_generado:
        archivos.append("apuntes_argos.docx")

    return {
        "flashcards": len(flashcards_unicas),
        "preguntas": len(preguntas),
        "archivos": archivos,
    }


def _generar_docx(carpeta: Path, datos: dict, preguntas: list, flashcards: list) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return False

    doc = Document()
    estilos = doc.styles
    estilos["Normal"].font.name = "Aptos"
    estilos["Normal"].font.size = Pt(10.5)

    doc.add_heading(datos.get("titulo", "Clase"), 0)
    doc.add_paragraph(f"Materia: {datos.get('materia', '')}")
    doc.add_paragraph("Documento generado automáticamente a partir de la transcripción. Requiere revisión.")

    doc.add_heading("Conceptos dominantes", level=1)
    doc.add_paragraph(", ".join(datos.get("palabras_clave_globales", [])) or "No detectados.")

    doc.add_heading("Índice temporal", level=1)
    for bloque in datos.get("bloques", []):
        doc.add_paragraph(
            f"{bloque.get('inicio')}–{bloque.get('fin')} · {bloque.get('titulo')}",
            style="List Bullet",
        )

    doc.add_heading("Apuntes por bloques", level=1)
    for bloque in datos.get("bloques", []):
        doc.add_heading(f"{bloque.get('numero')}. {bloque.get('titulo')}", level=2)
        doc.add_paragraph(f"Referencia: {bloque.get('inicio')}–{bloque.get('fin')}")
        doc.add_heading("Resumen", level=3)
        doc.add_paragraph(bloque.get("resumen", "") or "Sin resumen automático.")
        doc.add_heading("Conceptos clave", level=3)
        doc.add_paragraph(", ".join(bloque.get("palabras_clave", [])) or "—")
        if bloque.get("avisos_examen"):
            doc.add_heading("Avisos de examen", level=3)
            for aviso in bloque["avisos_examen"]:
                doc.add_paragraph(aviso, style="List Bullet")
        if bloque.get("preguntas"):
            doc.add_heading("Preguntas formuladas", level=3)
            for pregunta in bloque["preguntas"]:
                doc.add_paragraph(pregunta, style="List Bullet")
        doc.add_heading("Desarrollo limpio", level=3)
        doc.add_paragraph(bloque.get("texto", ""))

    doc.add_heading("Preguntas de repaso", level=1)
    for pregunta, minuto, bloque in preguntas:
        doc.add_paragraph(f"{pregunta} ({bloque}, {minuto})", style="List Number")

    doc.add_heading("Flashcards", level=1)
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    tabla.rows[0].cells[0].text = "Frente"
    tabla.rows[0].cells[1].text = "Dorso"
    tabla.rows[0].cells[2].text = "Minuto"
    for frente, dorso, minuto in flashcards:
        celdas = tabla.add_row().cells
        celdas[0].text = frente
        celdas[1].text = dorso
        celdas[2].text = minuto

    doc.save(carpeta / "apuntes_argos.docx")
    return True


def _frase_relevante(texto: str, palabra: str) -> str:
    frases = [x.strip() for x in re.split(r"(?<=[.!?])\s+|\n+", texto) if x.strip()]
    candidatas = [x for x in frases if palabra.casefold() in x.casefold()]
    if not candidatas:
        return ""
    mejor = max(candidatas, key=lambda x: min(len(x), 500))
    return mejor[:600]
